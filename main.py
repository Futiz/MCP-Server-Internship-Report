"""
Agent MCP Système de Fichiers

Un serveur Model Context Protocol qui fournit des capacités de
navigation et de recherche dans le système de fichiers.
Offre des outils pour lister les fichiers, lire le contenu
et rechercher dans les fichiers et répertoires.
"""

import fnmatch
import re
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
from mcp.server.fastmcp import FastMCP
import markdown
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls

mcp = FastMCP("FileSystem Agent")


@dataclass
class FileInfo:
    """Informations sur un fichier ou répertoire"""

    name: str
    path: str
    size: int
    is_directory: bool
    modified_time: float
    permissions: str


@dataclass
class SearchResult:
    """Résultat de recherche avec chemin du fichier, numéro de ligne et contenu correspondant"""

    file_path: str
    line_number: int
    line_content: str
    match_context: str


@mcp.tool()
def list_files(
    directory_path: str = ".", include_hidden: bool = False, pattern: str = "*"
) -> List[FileInfo]:
    """Liste les fichiers et répertoires dans le chemin spécifié.

    Args:
        directory_path: Chemin du répertoire à lister (défaut: répertoire courant)
        include_hidden: Inclure les fichiers/répertoires cachés (défaut: False)
        pattern: Motif glob pour filtrer les fichiers (défaut: "*" pour tous les fichiers)
    """
    try:
        path = Path(directory_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"Directory '{directory_path}' does not exist")
        if not path.is_dir():
            raise NotADirectoryError(f"'{directory_path}' is not a directory")

        files = []
        for item in path.iterdir():
            if not include_hidden and item.name.startswith("."):
                continue

            if not fnmatch.fnmatch(item.name, pattern):
                continue

            try:
                stat = item.stat()
                file_info = FileInfo(
                    name=item.name,
                    path=str(item),
                    size=stat.st_size,
                    is_directory=item.is_dir(),
                    modified_time=stat.st_mtime,
                    permissions=oct(stat.st_mode)[-3:],
                )
                files.append(file_info)
            except (OSError, PermissionError):
                continue

        return sorted(files, key=lambda x: (not x.is_directory, x.name.lower()))
    except Exception as e:
        raise RuntimeError(f"Error listing directory: {str(e)}") from e


def _is_text_file(file_path: Path) -> bool:
    """Détermine si un fichier est un fichier texte en lisant les premiers octets."""
    try:
        with open(file_path, "rb") as f:
            chunk = f.read(1024)
            if not chunk:
                return True  # Fichier vide = texte

            # Vérifier la présence de caractères null
            # (indicateur de fichier binaire)
            if b"\x00" in chunk:
                return False

            # Essayer de décoder en UTF-8
            try:
                chunk.decode("utf-8")
                return True
            except UnicodeDecodeError:
                pass

            # Essayer d'autres encodages courants
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    chunk.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue

            return False
    except (OSError, IOError, UnicodeDecodeError):
        return False


@mcp.tool()
def read_file(
    file_path: str, max_lines: int = 1000, encoding: str = "utf-8"
) -> Dict[str, Any]:
    """Lit le contenu d'un fichier.

    Args:
        file_path: Chemin du fichier à lire
        max_lines: Nombre maximum de lignes à lire (défaut: 1000)
        encoding: Encodage du fichier (défaut: utf-8)
    """
    try:
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File '{file_path}' does not exist")
        if path.is_dir():
            raise IsADirectoryError(f"'{file_path}' is a directory, not a file")

        # Vérifier si c'est un fichier texte
        if not _is_text_file(path):
            stat = path.stat()
            return {
                "file_path": str(path),
                "lines": [],
                "total_lines": 0,
                "truncated": False,
                "size_bytes": stat.st_size,
                "encoding": "binary",
                "error": (
                    f"Fichier binaire détecté (extension: {path.suffix}). "
                    "Utilisez un outil spécialisé pour lire ce type de fichier."
                ),
            }

        # Essayer de lire avec l'encodage spécifié
        encodings_to_try = (
            [encoding]
            if encoding != "utf-8"
            else ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
        )

        for enc in encodings_to_try:
            try:
                with open(path, "r", encoding=enc) as f:
                    lines = []
                    for i, line in enumerate(f, 1):
                        if i > max_lines:
                            break
                        lines.append(line.rstrip("\n\r"))

                stat = path.stat()
                return {
                    "file_path": str(path),
                    "lines": lines,
                    "total_lines": len(lines),
                    "truncated": len(lines) == max_lines,
                    "size_bytes": stat.st_size,
                    "encoding": enc,
                }
            except UnicodeDecodeError as exc:
                if enc == encodings_to_try[-1]:  # Dernier encodage testé
                    raise RuntimeError(
                        f"Impossible de décoder le fichier avec les encodages "
                        f"testés: {encodings_to_try}"
                    ) from exc
                continue

    except Exception as e:
        raise RuntimeError(f"Erreur lors de la lecture du fichier: {str(e)}") from e


@mcp.tool()
def search_content(
    search_term: str,
    directory_path: str = ".",
    file_pattern: str = "*",
    case_sensitive: bool = False,
    max_results: int = 100,
) -> List[SearchResult]:
    """Recherche du contenu textuel dans les fichiers.

    Args:
        search_term: Texte à rechercher
        directory_path: Répertoire dans lequel rechercher (défaut: répertoire courant)
        file_pattern: Motif de fichier à correspondre (défaut: "*" pour tous les fichiers)
        case_sensitive: Si la recherche doit être sensible à la casse (défaut: False)
        max_results: Nombre maximum de résultats à retourner (défaut: 100)
    """
    try:
        path = Path(directory_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"Directory '{directory_path}' does not exist")

        results = []
        search_lower = search_term.lower() if not case_sensitive else search_term

        def search_file(file_path: Path):
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    for line_num, line in enumerate(f, 1):
                        line_to_search = line if case_sensitive else line.lower()
                        if search_lower in line_to_search:
                            context_start = max(0, line_to_search.find(search_lower) - 20)
                            context_end = min(
                                len(line), context_start + len(search_lower) + 40
                            )
                            context = line[context_start:context_end].strip()

                            results.append(
                                SearchResult(
                                    file_path=str(file_path),
                                    line_number=line_num,
                                    line_content=line.strip(),
                                    match_context=context,
                                )
                            )

                            if len(results) >= max_results:
                                return False
                return True
            except (UnicodeDecodeError, PermissionError, OSError):
                return True

        if path.is_file():
            search_file(path)
        else:
            for file_path in path.rglob(file_pattern):
                if file_path.is_file():
                    if not search_file(file_path):
                        break

        return results
    except Exception as e:
        raise RuntimeError(f"Error searching content: {str(e)}") from e


@mcp.tool()
def get_binary_file_info(file_path: str) -> Dict[str, Any]:
    """Obtient des informations sur un fichier binaire.

    Args:
        file_path: Chemin du fichier binaire à analyser
    """
    try:
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File '{file_path}' does not exist")
        if path.is_dir():
            raise IsADirectoryError(f"'{file_path}' is a directory, not a file")

        stat = path.stat()

        # Lire les premiers octets pour identifier le type
        magic_bytes = b""
        try:
            with open(path, "rb") as f:
                magic_bytes = f.read(16)
        except (OSError, IOError):
            pass

        # Détection basique du type de fichier
        file_type = "Unknown binary"
        if magic_bytes.startswith(b"%PDF"):
            file_type = "PDF Document"
        elif magic_bytes.startswith(b"\x89PNG"):
            file_type = "PNG Image"
        elif magic_bytes.startswith(b"\xff\xd8\xff"):
            file_type = "JPEG Image"
        elif magic_bytes.startswith(b"GIF8"):
            file_type = "GIF Image"
        elif magic_bytes.startswith(b"PK"):
            file_type = "ZIP Archive (ou fichier Office)"
        elif magic_bytes.startswith(b"\x7fELF"):
            file_type = "Executable Linux/Unix"
        elif magic_bytes.startswith(b"MZ"):
            file_type = "Executable Windows"

        return {
            "file_path": str(path),
            "name": path.name,
            "size_bytes": stat.st_size,
            "modified_time": stat.st_mtime,
            "permissions": oct(stat.st_mode)[-3:],
            "extension": path.suffix,
            "detected_type": file_type,
            "magic_bytes_hex": magic_bytes.hex() if magic_bytes else "",
            "is_binary": True,
        }
    except Exception as e:
        raise RuntimeError(f"Erreur lors de l'analyse du fichier binaire: {str(e)}") from e


@mcp.tool()
def read_binary_file(
    file_path: str, max_bytes: int = 1024, offset: int = 0
) -> Dict[str, Any]:
    """Lit le contenu brut d'un fichier binaire.

    Args:
        file_path: Chemin du fichier binaire à lire
        max_bytes: Nombre maximum d'octets à lire (défaut: 1024)
        offset: Position de départ dans le fichier (défaut: 0)
    """
    try:
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File '{file_path}' does not exist")
        if path.is_dir():
            raise IsADirectoryError(f"'{file_path}' is a directory, not a file")

        stat = path.stat()

        with open(path, "rb") as f:
            if offset > 0:
                f.seek(offset)

            content = f.read(max_bytes)

            # Convertir en représentations lisibles
            hex_content = content.hex()

            # Essayer de créer une représentation ASCII (avec . pour les non-imprimables)
            ascii_repr = ""
            for byte in content:
                if 32 <= byte <= 126:  # Caractères imprimables ASCII
                    ascii_repr += chr(byte)
                else:
                    ascii_repr += "."

            return {
                "file_path": str(path),
                "total_size_bytes": stat.st_size,
                "bytes_read": len(content),
                "offset": offset,
                "hex_content": hex_content,
                "ascii_representation": ascii_repr,
                "truncated": (
                    len(content) == max_bytes and offset + max_bytes < stat.st_size
                ),
            }
    except Exception as e:
        raise RuntimeError(f"Erreur lors de la lecture du fichier binaire: {str(e)}") from e


@mcp.tool()
def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """Extrait le texte d'un fichier PDF (basique, sans dépendances externes)

    Args:
        file_path: Chemin du fichier PDF
    """
    try:
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File '{file_path}' does not exist")
        if path.suffix.lower() != ".pdf":
            raise ValueError(f"Le fichier '{file_path}' n'est pas un PDF")

        # Lecture basique du PDF pour extraire du texte visible
        text_content = []

        with open(path, "rb") as f:
            content = f.read()

            # Recherche simple de texte dans les streams PDF
            # Cette méthode est très basique et ne fonctionne que pour certains PDFs

            # Chercher les objets de texte dans le PDF
            text_objects = re.findall(rb"BT\s.*?ET", content, re.DOTALL)

            for obj in text_objects:
                # Extraire les chaînes entre parenthèses ou crochets
                strings = re.findall(rb"[\(\[]([^)\]]*?)[\)\]]", obj)
                for s in strings:
                    try:
                        decoded = s.decode("latin-1", errors="ignore")
                        if decoded.strip() and len(decoded) > 2:
                            text_content.append(decoded.strip())
                    except (OSError, IOError, UnicodeDecodeError):
                        continue

        stat = path.stat()
        return {
            "file_path": str(path),
            "size_bytes": stat.st_size,
            "extracted_text_lines": text_content,
            "total_text_fragments": len(text_content),
            "warning": (
                "Extraction basique - certains PDFs peuvent nécessiter "
                "des outils spécialisés comme PyPDF2"
            ),
        }
    except Exception as e:
        raise RuntimeError(f"Erreur lors de l'extraction de texte du PDF: {str(e)}") from e


@mcp.resource("directory://{path}")
def get_directory_info(path: str) -> Dict[str, Any]:
    """Obtient des informations détaillées sur un répertoire"""
    try:
        dir_path = Path(path).resolve()
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory '{path}' does not exist")
        if not dir_path.is_dir():
            raise NotADirectoryError(f"'{path}' is not a directory")

        files = list_files(str(dir_path), include_hidden=True)
        total_size = sum(f.size for f in files if not f.is_directory)

        return {
            "path": str(dir_path),
            "total_files": len([f for f in files if not f.is_directory]),
            "total_directories": len([f for f in files if f.is_directory]),
            "total_size_bytes": total_size,
            "files": [f.__dict__ for f in files],
        }
    except Exception as e:
        raise RuntimeError(f"Error getting directory info: {str(e)}") from e


@mcp.resource("file://{path}")
def get_file_info(path: str) -> Dict[str, Any]:
    """Obtient des informations détaillées sur un fichier"""
    try:
        file_path = Path(path).resolve()
        if not file_path.exists():
            raise FileNotFoundError(f"File '{path}' does not exist")
        if file_path.is_dir():
            raise IsADirectoryError(f"'{path}' is a directory, not a file")

        stat = file_path.stat()

        is_text = True
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(
                    1024
                )  # Lit les 1024 premiers caractères pour vérifier si c'est un fichier texte
        except UnicodeDecodeError:
            is_text = False

        return {
            "path": str(file_path),
            "name": file_path.name,
            "size_bytes": stat.st_size,
            "modified_time": stat.st_mtime,
            "permissions": oct(stat.st_mode)[-3:],
            "is_text_file": is_text,
            "extension": file_path.suffix,
            "parent_directory": str(file_path.parent),
        }
    except Exception as e:
        raise RuntimeError(f"Error getting file info: {str(e)}") from e


# Configuration centralisée des templates
TEMPLATE_CONFIGS = {
        "présentation de l'entreprise": """
STRUCTURE RECOMMANDÉE:
1. Historique et évolution de l'entreprise
2. Secteur d'activité et positionnement concurrentiel
3. Organisation et structure hiérarchique
4. Chiffres clés (effectifs, CA, implantations géographiques)
5. Valeurs et culture d'entreprise

POINTS CLÉS À DÉVELOPPER:
- Contexte économique et enjeux du secteur
- Innovation et stratégie de développement
- Positionnement sur le marché
- Relations avec les partenaires/clients principaux

INFORMATIONS À COLLECTER:
- Date de création, fondateurs, étapes clés
- Domaines d'expertise et technologies maîtrisées
- Organigramme et répartition des équipes
- Projets phares et références clients
""",
        
        "service d'accueil": """
STRUCTURE RECOMMANDÉE:
1. Position du service dans l'organigramme général
2. Missions et responsabilités du service
3. Composition de l'équipe et profils
4. Interactions avec les autres services
5. Enjeux spécifiques liés au stage

POINTS CLÉS À DÉVELOPPER:
- Rôle stratégique du service dans l'entreprise
- Méthodologies de travail et processus
- Technologies et outils utilisés
- Défis actuels et projets en cours

INFORMATIONS À COLLECTER:
- Nom et fonction du maître de stage
- Expertise technique de l'équipe
- Budget et ressources allouées
- Objectifs à court et moyen terme
""",

        "mission": """
STRUCTURE RECOMMANDÉE:
1. Problématique technique ou scientifique
2. Contexte et enjeux pour l'entreprise
3. Objectifs fixés et livrables attendus
4. Périmètre et contraintes du projet
5. Planning prévisionnel vs réalisé

POINTS CLÉS À DÉVELOPPER:
- Analyse du besoin initial
- Complexité technique et défis identifiés
- Ressources mises à disposition
- Critères d'évaluation du succès

INFORMATIONS À COLLECTER:
- Cahier des charges détaillé
- Acteurs impliqués dans le projet
- Budget et délais impartis
- Risques identifiés en amont
""",

        "état de l'art": """
STRUCTURE RECOMMANDÉE:
1. Technologies existantes sur le marché
2. Solutions internes déjà en place
3. Benchmarking des approches concurrentes
4. Avantages/inconvénients de chaque solution
5. Positionnement de l'approche choisie

POINTS CLÉS À DÉVELOPPER:
- Analyse comparative rigoureuse
- Critères de sélection technique
- Évolution technologique du domaine
- Retour d'expérience d'autres projets

INFORMATIONS À COLLECTER:
- Documentation technique des solutions
- Études de marché et rapports sectoriels  
- Retours utilisateurs et cas d'usage
- Coûts de mise en œuvre et maintenance
""",

        "méthodologie": """
STRUCTURE RECOMMANDÉE:
1. Approche méthodologique adoptée
2. Phases du projet et jalons
3. Outils et technologies sélectionnés
4. Justification des choix techniques
5. Métriques de suivi et validation

POINTS CLÉS À DÉVELOPPER:
- Adéquation méthode/problématique
- Processus de prise de décision
- Gestion des risques et plan B
- Adaptations en cours de projet

INFORMATIONS À COLLECTER:
- Méthodologies standards du secteur
- Contraintes techniques et organisationnelles
- Formation reçue sur les outils
- Retours d'expérience équipe
""",

        "difficultés": """
STRUCTURE RECOMMANDÉE:
1. Difficultés techniques rencontrées
2. Problèmes organisationnels ou humains
3. Contraintes temporelles ou budgétaires
4. Solutions palliatives mises en place
5. Leçons apprises pour l'avenir

POINTS CLÉS À DÉVELOPPER:
- Analyse des causes profondes
- Impact sur le planning et les objectifs
- Créativité dans les solutions trouvées
- Capacité d'adaptation et de résilience

INFORMATIONS À COLLECTER:
- Chronologie des problèmes
- Ressources mobilisées pour les résoudre
- Aide reçue de l'équipe/hiérarchie
- Amélioration des processus suite aux difficultés
""",

        "résultats": """
STRUCTURE RECOMMANDÉE:
1. Présentation des livrables finaux
2. Métriques et indicateurs de performance
3. Comparaison objectifs vs réalisations
4. Validation par les parties prenantes
5. Documentation et transfert de compétences

POINTS CLÉS À DÉVELOPPER:
- Démonstration concrète des résultats
- Analyse quantitative et qualitative
- Valeur ajoutée pour l'entreprise
- Perspectives d'évolution du projet

INFORMATIONS À COLLECTER:
- Captures d'écran, prototypes, démos
- Métriques avant/après implementation
- Feedback des utilisateurs finaux
- ROI estimé ou gains mesurables
""",

        "analyse critique": """
STRUCTURE RECOMMANDÉE:
1. Forces et faiblesses de la solution
2. Comparaison avec l'état de l'art
3. Limites identifiées et améliorations possibles
4. Pertinence par rapport aux objectifs initiaux
5. Recommandations pour la suite

POINTS CLÉS À DÉVELOPPER:
- Objectivité dans l'évaluation
- Vision critique et constructive
- Prise de recul sur les choix effectués
- Maturité dans l'analyse technique

INFORMATIONS À COLLECTER:
- Tests de performance détaillés
- Comparatifs avec solutions existantes
- Feedback des experts du domaine
- Évolutions technologiques à venir
""",

        "retombées": """
STRUCTURE RECOMMANDÉE:
1. Impact immédiat sur les processus
2. Économies ou gains de productivité
3. Amélioration de la performance technique
4. Perspectives de déploiement élargi
5. Contribution à la stratégie d'innovation

POINTS CLÉS À DÉVELOPPER:
- Quantification des bénéfices
- Adoption par les équipes métier
- Scalabilité de la solution
- Positionnement concurrentiel renforcé

INFORMATIONS À COLLECTER:
- Métriques business avant/après
- Retours des équipes utilisatrices
- Plans de déploiement futurs
- Valorisation potentielle de la propriété intellectuelle
""",

        "bilan personnel": """
STRUCTURE RECOMMANDÉE:
1. Compétences techniques acquises
2. Développement des soft skills
3. Compréhension du monde de l'entreprise
4. Réseau professionnel constitué
5. Impact sur le projet professionnel

POINTS CLÉS À DÉVELOPPER:
- Evolution personnelle mesurable
- Confrontation théorie/pratique
- Autonomie progressivement gagnée
- Capacité de remise en question

INFORMATIONS À COLLECTER:
- Auto-évaluation des compétences
- Retours du maître de stage
- Moments marquants du stage
- Liens avec le parcours de formation
""",

        "compétences acquises": """
STRUCTURE RECOMMANDÉE:
1. Compétences techniques spécialisées
2. Maîtrise des outils professionnels
3. Méthodologies de travail intégrées
4. Compétences transversales développées
5. Certification ou formation complémentaire

POINTS CLÉS À DÉVELOPPER:
- Concrétisation par des exemples précis
- Niveau de maîtrise atteint
- Transférabilité vers d'autres contextes
- Valeur ajoutée sur le CV

INFORMATIONS À COLLECTER:
- Portfolio des réalisations techniques
- Certifications obtenues pendant le stage
- Formations suivies en parallèle
- Feedback des collègues sur les progrès
""",

        "perspectives professionnelles": """
STRUCTURE RECOMMANDÉE:
1. Clarification du projet professionnel
2. Secteurs d'activité d'intérêt
3. Types de postes envisagés
4. Compétences à développer davantage
5. Suite de parcours (études, emploi)

POINTS CLÉS À DÉVELOPPER:
- Cohérence avec les aspirations initiales
- Influence du stage sur les choix futurs
- Réalisme des perspectives
- Plan de développement personnel

INFORMATIONS À COLLECTER:
- Discussions avec le tuteur entreprise
- Rencontres avec d'autres professionnels
- Analyse du marché de l'emploi
- Opportunités identifiées dans l'entreprise
"""
}

# Template générique par défaut
DEFAULT_TEMPLATE = """
STRUCTURE GÉNÉRIQUE:
1. Introduction du sujet
2. Développement des points principaux
3. Analyse et réflexion critique
4. Conclusion et ouvertures

POINTS CLÉS À DÉVELOPPER:
- Contextualisation par rapport au stage
- Exemples concrets et détaillés
- Liens avec les objectifs du rapport
- Perspective personnelle et professionnelle
"""


def _get_section_template(section: str) -> str:
    """Retourne le template spécifique pour chaque section du rapport de stage"""
    
    # Normaliser le nom de la section pour la comparaison
    section_norm = section.lower().strip()
    
    # Recherche par correspondance partielle
    for key, template in TEMPLATE_CONFIGS.items():
        if key in section_norm or any(word in section_norm for word in key.split()):
            return template
    
    return DEFAULT_TEMPLATE


@mcp.prompt()
def generate_internship_report(section: str, content_details: str) -> str:
    """Génère une section du rapport de stage avec template et directives intégrées

    Args:
        section: Section à générer (ex: "introduction", "contexte", "missions", etc.)
        content_details: Détails spécifiques à inclure dans cette section
    """
    
    # Obtenir le template spécifique à la section
    section_template = _get_section_template(section)

    # Directives de rédaction intégrées directement dans le prompt
    directives = """DIRECTIVES DE RÉDACTION OBLIGATOIRES POUR RAPPORT DE STAGE:

STYLE ET STRUCTURE:
- Utiliser un style académique professionnel mais accessible
- Structurer avec des paragraphes courts et aérés (3-5 phrases max)
- Employer la première personne avec parcimonie ("j'ai observé", "j'ai contribué")
- Privilégier les phrases actives et concrètes
- Utiliser des connecteurs logiques pour fluidifier le texte

CONTENU ET APPROCHE:
- Contexturaliser chaque section par rapport à l'entreprise et au secteur
- Intégrer des données chiffrées et exemples concrets quand possible
- Montrer l'évolution/progression de vos compétences
- Faire le lien entre théorie (formation) et pratique (stage)
- Adopter une posture réflexive et analytique

EXIGENCES TECHNIQUES:
- Respecter la terminologie professionnelle du domaine
- Citer les outils, méthodologies et technologies utilisées
- Expliquer les enjeux business/techniques sans jargon excessif
- Inclure une dimension prospective (apprentissages, évolutions)

QUALITÉ RÉDACTIONNELLE:
- Vérifier la cohérence des temps verbaux
- Éviter les répétitions et redondances
- Utiliser un vocabulaire varié et précis
- Soigner les transitions entre idées"""

    return f"""{directives}

==================================================
TEMPLATE SPÉCIFIQUE POUR: {section.upper()}
{section_template}

==================================================
SECTION À GÉNÉRER: {section.upper()}
DÉTAILS SPÉCIFIQUES: {content_details}

En respectant scrupuleusement les directives ci-dessus ET en suivant le template spécifique, génère maintenant cette section du rapport de stage de manière professionnelle et structurée."""


@mcp.tool()
def list_available_templates() -> List[str]:
    """Liste tous les templates de sections disponibles pour le rapport de stage"""
    # Récupérer les clés depuis _get_section_template pour éviter les doublons
    templates = {
        "présentation de l'entreprise": None,
        "service d'accueil": None, 
        "mission": None,
        "état de l'art": None,
        "méthodologie": None, 
        "difficultés": None,
        "résultats": None,
        "analyse critique": None,
        "retombées": None,
        "bilan personnel": None,
        "compétences acquises": None,
        "perspectives professionnelles": None
    }
    
    return list(templates.keys())


@mcp.tool()  
def get_section_template_preview(section: str) -> str:
    """Affiche le template d'une section spécifique sans générer de contenu
    
    Args:
        section: Nom de la section (ex: "mission", "résultats", etc.)
    """
    template = _get_section_template(section)
    
    return f"""📋 TEMPLATE POUR LA SECTION: {section.upper()}
{'-' * 50}
{template}
{'-' * 50}

💡 Pour utiliser ce template, appelez generate_internship_report() avec:
- section: "{section}" 
- content_details: [vos détails spécifiques]
"""


@mcp.tool()
def update_report_section(
    section: str,
    markdown_file: str,
    new_section_content: str,
    create_if_missing: bool = True,
) -> str:
    """Met à jour une section spécifique dans un fichier Markdown de rapport de stage

    Args:
        section: Section à mettre à jour (ex: "introduction", "contexte", "missions", etc.)
        content_details: Détails spécifiques à inclure dans cette section
        markdown_file: Chemin vers le fichier .md du rapport
        new_section_content: Le nouveau contenu de la section (généré selon les directives)
        create_if_missing: Créer la section si elle n'existe pas (défaut: True)
    """

    try:
        # Vérifier si le fichier existe et le lire
        file_path = Path(markdown_file).resolve()
        current_content = _read_file_safe(markdown_file)
        
        if not current_content and not create_if_missing:
            return f"ERREUR: Le fichier '{markdown_file}' n'existe pas."

        # Helpers: normaliser un titre et retirer un éventuel header du nouveau contenu
        def _normalize_title(s: str) -> str:
            s_norm = s.strip().lower()
            s_norm = re.sub(r"\s+", " ", s_norm)
            # Retire une numérotation de début (ex: "2.", "2)", "II.") simple
            s_norm = re.sub(r"^(\d+|[ivxlcdm]+)[\.)\-:]\s*", "", s_norm)
            return s_norm

        def _strip_leading_header(content: str) -> str:
            lines = content.lstrip().splitlines()
            if lines and re.match(r"^#{1,6}\s+", lines[0]):
                return "\n".join(lines[1:]).lstrip()
            return content.strip()

        target_norm = _normalize_title(section)
        updated_content = current_content

        # Détecter tous les headers markdown et comparer les titres normalisés
        header_iter = list(re.finditer(r"^(#{1,6})\s+(.+)$", current_content, re.MULTILINE))
        section_found = False

        for idx, m in enumerate(header_iter):
            header_title = m.group(2).strip()
            header_title_norm = _normalize_title(header_title)

            if header_title_norm == target_norm:
                # Déterminer la fin de la section (avant le prochain header)
                start = m.start()
                end = header_iter[idx + 1].start() if idx + 1 < len(header_iter) else len(current_content)

                existing_header_line = current_content[start: current_content.find("\n", start) if current_content.find("\n", start) != -1 else end]
                body_replacement = _strip_leading_header(new_section_content)

                replacement = existing_header_line + "\n\n" + body_replacement.strip() + "\n\n"
                updated_content = current_content[:start] + replacement + current_content[end:]
                section_found = True
                break

        # Si la section n'a pas été trouvée, on ajoute une nouvelle section en fin de fichier
        if not section_found:
            section_header = f"## {section.title()}"
            body_replacement = _strip_leading_header(new_section_content)
            formatted_new_content = f"{section_header}\n\n{body_replacement}\n\n"
            if current_content and not current_content.endswith("\n\n"):
                updated_content = current_content.rstrip() + "\n\n" + formatted_new_content
            else:
                updated_content = current_content + formatted_new_content

        # Écrire le contenu modifié
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(updated_content)

        # Retourner le statut de la modification
        action = "modifiée" if section_found else "ajoutée"
        return f"""✅ MODIFICATION EFFECTUÉE

Fichier: {markdown_file}
Section: {section.title()} ({action})
Taille du nouveau contenu: {len(new_section_content)} caractères

Le fichier a été mis à jour avec succès."""

    except Exception as e:
        return f"Erreur lors de la modification du fichier: {str(e)}"


def _read_file_safe(file_path: str) -> str:
    """Lit un fichier de manière sécurisée avec gestion d'erreur"""
    try:
        path = Path(file_path).resolve()
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
    except Exception:
        pass
    return ""


@mcp.tool()
def get_report_section_instructions(
    section: str, content_details: str, markdown_file: str
) -> str:
    """Obtient les instructions pour rédiger une section 
    de rapport (à utiliser avant update_report_section)

    Args:
        section: Section à rédiger (ex: "introduction", "contexte", "missions", etc.)
        content_details: Détails spécifiques à inclure dans cette section
        markdown_file: Chemin vers le fichier .md du rapport (pour contexte)
    """

    # Utiliser le prompt intégré pour générer les instructions
    prompt_instructions = generate_internship_report(section, content_details)

    # Lire le fichier actuel pour contexte
    current_content = _read_file_safe(markdown_file)

    result = f"""📝 INSTRUCTIONS POUR RÉDACTION DE SECTION

FICHIER CIBLE: {markdown_file}
SECTION: {section.upper()}

CONTENU ACTUEL DU FICHIER:
{'-' * 50}
{current_content if current_content else '[Fichier vide ou inexistant]'}
{'-' * 50}

{prompt_instructions}

⚠️ ÉTAPES SUIVANTES:
1. Rédigez le contenu de la section selon les directives ci-dessus
2. Utilisez ensuite l'outil 'update_report_section' avec le contenu rédigé"""

    return result


@mcp.tool()
def validate_report_structure(markdown_file: str) -> Dict[str, Any]:
    """Valide la structure du rapport de stage et identifie les sections manquantes
    
    Args:
        markdown_file: Chemin vers le fichier .md du rapport
    """
    try:
        current_content = _read_file_safe(markdown_file)
        
        if not current_content:
            return {
                "status": "error",
                "message": f"Fichier '{markdown_file}' vide ou inexistant",
                "sections_found": [],
                "sections_missing": list(TEMPLATE_CONFIGS.keys()),
                "recommendations": ["Commencer par créer le fichier et ajouter des sections"]
            }
        
        # Extraire tous les headers du document
        headers = re.findall(r"^(#{1,6})\s+(.+)$", current_content, re.MULTILINE)
        sections_found = []
        
        for level, title in headers:
            title_norm = title.strip().lower()
            title_norm = re.sub(r"^(\d+|[ivxlcdm]+)[\.)\-:]\s*", "", title_norm)
            sections_found.append({
                "level": len(level),
                "title": title.strip(),
                "normalized": title_norm
            })
        
        # Identifier les sections du rapport présentes
        template_keys = set(TEMPLATE_CONFIGS.keys())
        found_keys = set()
        
        for section in sections_found:
            for key in template_keys:
                if key in section["normalized"] or any(word in section["normalized"] for word in key.split()):
                    found_keys.add(key)
                    break
        
        sections_missing = template_keys - found_keys
        completion_rate = len(found_keys) / len(template_keys) * 100
        
        # Générer des recommandations
        recommendations = []
        if completion_rate < 50:
            recommendations.append("Structure de base incomplète - ajouter les sections principales")
        if "présentation de l'entreprise" not in found_keys:
            recommendations.append("Ajouter la présentation de l'entreprise (section fondamentale)")
        if "mission" not in found_keys:
            recommendations.append("Décrire la mission principale du stage")
        if "bilan personnel" not in found_keys:
            recommendations.append("Inclure un bilan personnel et retour d'expérience")
            
        return {
            "status": "success",
            "completion_rate": round(completion_rate, 1),
            "total_sections": len(sections_found),
            "sections_found": [s["title"] for s in sections_found],
            "template_sections_present": sorted(list(found_keys)),
            "sections_missing": sorted(list(sections_missing)),
            "recommendations": recommendations,
            "file_stats": {
                "word_count": len(current_content.split()),
                "char_count": len(current_content),
                "line_count": len(current_content.splitlines())
            }
        }
        
    except Exception as e:
        return {
            "status": "error", 
            "message": f"Erreur lors de la validation: {str(e)}",
            "sections_found": [],
            "sections_missing": [],
            "recommendations": []
        }


@mcp.tool()
def convert_markdown_to_docx(
    markdown_file: str,
    output_file: str = None,
    title: str = "Rapport de Stage",
    author: str = "",
    font_name: str = "Times New Roman",
    font_size: int = 12,
    margin_cm: float = 2.5,
    line_spacing: float = 1.5,
    language: str = "en-US",
    add_page_numbers: bool = True,
    decimal_numbering: bool = True
) -> Dict[str, Any]:
    """Convertit un fichier Markdown en DOCX avec formatage académique professionnel
    
    Args:
        markdown_file: Chemin vers le fichier .md source
        output_file: Chemin de sortie .docx (optionnel, généré automatiquement si absent)
        title: Titre du document (défaut: "Rapport de Stage")
        author: Nom de l'auteur
        font_name: Police à utiliser (défaut: Times New Roman - norme académique)
        font_size: Taille de police en points (défaut: 12pt - norme académique)
        margin_cm: Marges en centimètres (défaut: 2.5cm - norme académique)
        line_spacing: Interligne (défaut: 1.5 - norme académique)
        language: Langue du document (défaut: en-US)
        add_page_numbers: Ajouter numérotation des pages (défaut: True)
        decimal_numbering: Utiliser numérotation décimale pour les titres (défaut: True)
    """
    try:
        # Lire le fichier markdown
        markdown_content = _read_file_safe(markdown_file)
        if not markdown_content:
            return {
                "status": "error",
                "message": f"Impossible de lire le fichier '{markdown_file}'"
            }
        
        # Déterminer le fichier de sortie
        if not output_file:
            md_path = Path(markdown_file)
            output_file = str(md_path.with_suffix('.docx'))
        
        # Créer le document Word
        doc = Document()
        
        # Configuration du document selon les normes académiques
        _setup_academic_document(doc, font_name, font_size, margin_cm, line_spacing, language, add_page_numbers)
        
        # Parser le markdown et convertir avec numérotation
        _convert_markdown_content(doc, markdown_content, font_name, font_size, line_spacing, decimal_numbering)
        
        # Ajouter métadonnées
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Sauvegarder le document
        doc.save(output_file)
        
        # Statistiques du document créé
        output_path = Path(output_file)
        file_size = output_path.stat().st_size if output_path.exists() else 0
        
        return {
            "status": "success",
            "input_file": markdown_file,
            "output_file": output_file,
            "file_size_bytes": file_size,
            "format_settings": {
                "font": f"{font_name} {font_size}pt",
                "margins": f"{margin_cm} cm",
                "line_spacing": f"{line_spacing}",
                "language": language,
                "page_numbers": add_page_numbers,
                "decimal_numbering": decimal_numbering,
                "title": title,
                "author": author
            },
            "message": f"Document DOCX créé avec succès : {output_file}"
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Erreur lors de la conversion: {str(e)}"
        }


def _setup_academic_document(doc: Document, font_name: str, font_size: int, margin_cm: float, line_spacing: float, language: str, add_page_numbers: bool):
    """Configure le document selon les normes académiques"""
    # Configuration des sections et marges (2.5cm minimum)
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    
    # Taille du papier A4
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    
    # Ajout de la numérotation des pages si demandé
    if add_page_numbers:
        _add_page_numbers(doc, section)
    
    # Configuration des styles
    _setup_academic_styles(doc, font_name, font_size, line_spacing, language)


def _add_page_numbers(doc: Document, section):
    """Ajoute la numérotation des pages"""
    try:
        # Créer le footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le numéro de page
        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        
        # Code XML pour insérer le numéro de page
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
    except Exception:
        pass  # Si l'ajout automatique échoue, continuer sans numérotation


def _setup_academic_styles(doc: Document, font_name: str, font_size: int, line_spacing: float, language: str):
    """Configure les styles selon les normes académiques"""
    styles = doc.styles
    
    # Style normal avec interligne 1.5
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = font_name
    normal_font.size = Pt(font_size)
    
    # Configuration de l'interligne
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_paragraph_format.line_spacing = line_spacing
    normal_paragraph_format.space_after = Pt(6)
    
    # Styles pour les titres avec numérotation
    heading_sizes = [18, 16, 14, 13, 12, 12]  # Tailles décroissantes pour H1-H6
    
    for i in range(1, 7):
        try:
            heading_style = styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = font_name
            heading_font.size = Pt(heading_sizes[i-1])
            heading_font.bold = True
            
            # Espacement des titres
            heading_format = heading_style.paragraph_format
            heading_format.space_before = Pt(12)
            heading_format.space_after = Pt(6)
            heading_format.keep_with_next = True
        except KeyError:
            pass


def _convert_markdown_content(doc: Document, content: str, font_name: str, font_size: int, line_spacing: float, decimal_numbering: bool):
    """Convertit le contenu markdown avec gestion de la numérotation décimale"""
    lines = content.split('\n')
    heading_counters = [0, 0, 0, 0, 0, 0]  # Compteurs pour H1-H6
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        line_stripped = line.strip()
        
        # Gestion des blocs de code
        if line_stripped.startswith('```'):
            if in_code_block:
                _add_academic_code_block(doc, '\n'.join(code_block_content))
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Gestion des titres avec numérotation décimale
        if line_stripped.startswith('#'):
            level = len(line_stripped) - len(line_stripped.lstrip('#'))
            title_text = line_stripped.lstrip('# ').strip()
            
            if decimal_numbering:
                title_text = _add_decimal_numbering(title_text, level, heading_counters)
            
            _add_academic_heading(doc, title_text, level)
        
        # Listes
        elif line_stripped.startswith(('- ', '* ', '+ ')):
            item_text = line_stripped[2:].strip()
            _add_academic_list_item(doc, item_text, font_name, font_size)
        
        elif line_stripped.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            item_text = line_stripped.split('. ', 1)[1] if '. ' in line_stripped else line_stripped
            _add_academic_numbered_list_item(doc, item_text, font_name, font_size)
        
        # Paragraphes normaux
        elif line_stripped:
            _add_academic_paragraph(doc, line_stripped, font_name, font_size, line_spacing)
        
        # Lignes vides
        else:
            doc.add_paragraph()


def _add_decimal_numbering(title: str, level: int, counters: List[int]) -> str:
    """Ajoute la numérotation décimale aux titres"""
    # Incrémenter le compteur du niveau actuel
    counters[level - 1] += 1
    
    # Remettre à zéro les compteurs des niveaux inférieurs
    for i in range(level, 6):
        counters[i] = 0
    
    # Construire le numéro décimal
    number_parts = []
    for i in range(level):
        if counters[i] > 0:
            number_parts.append(str(counters[i]))
    
    if number_parts:
        decimal_number = '.'.join(number_parts)
        return f"{decimal_number}. {title}"
    
    return title


def _add_academic_heading(doc: Document, text: str, level: int):
    """Ajoute un titre avec formatage académique"""
    level = min(max(level, 1), 6)
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_academic_paragraph(doc: Document, text: str, font_name: str, font_size: int, line_spacing: float):
    """Ajoute un paragraphe avec formatage académique"""
    text = _process_inline_markdown(text)
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    
    # Interligne et espacement
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_after = Pt(6)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_academic_list_item(doc: Document, text: str, font_name: str, font_size: int):
    """Ajoute un élément de liste à puces avec formatage académique"""
    text = _process_inline_markdown(text)
    paragraph = doc.add_paragraph(text, style='List Bullet')
    
    # Appliquer le formatage
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _add_academic_numbered_list_item(doc: Document, text: str, font_name: str, font_size: int):
    """Ajoute un élément de liste numérotée avec formatage académique"""
    text = _process_inline_markdown(text)
    paragraph = doc.add_paragraph(text, style='List Number')
    
    # Appliquer le formatage
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _add_academic_code_block(doc: Document, code: str):
    """Ajoute un bloc de code avec formatage académique"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'  # Police monospace standard
    run.font.size = Pt(10)
    
    # Style pour le code avec bordure
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(1)
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)


def _process_inline_markdown(text: str) -> str:
    """Traite le markdown inline basique"""
    # Note: Pour respecter les normes académiques, on garde le formatage simple
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Gras -> texte normal
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italique -> texte normal  
    text = re.sub(r'`(.*?)`', r'\1', text)        # Code inline -> texte normal
    return text
